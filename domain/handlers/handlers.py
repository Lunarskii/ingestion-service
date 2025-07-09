from abc import (
    ABC,
    abstractmethod,
)
from pathlib import Path
from typing import Literal
from datetime import datetime
import logging

from pypdf import (
    PdfReader,
    DocumentInformation,
)
from bs4 import BeautifulSoup
from pydantic import BaseModel
import langdetect
import docx
import openpyxl

import utils


class ProcessingResult(BaseModel):
    status: str
    error_message: str
    path: str
    detected_language: str | None = None
    document_page_count: int | None = None
    author: str | None = None
    creation_date: datetime | None = None


class DocumentHandler(ABC):
    def __init__(self, *, dest_folder: str = "."):
        self.dest_folder: str = dest_folder
        self.status: Literal["success", "failed_processing"] = "success"
        self.error_message: str = ""
        self.path: str = ""

    def handle(
        self,
        file_path: str,
        *,
        dest_folder: str | None = None,
        file_name: str | None = None,
    ) -> ProcessingResult:
        dest_folder: Path = self._mkdir(dest_folder or self.dest_folder)
        file_name: str = file_name or file_path.split("/")[-1].split(".")[0] + ".txt"
        dest_file_path = str(dest_folder / file_name)

        try:
            logging.info(f"Starting new processing {file_path}")
            result: ProcessingResult = self._handle(file_path, dest_file_path)
        except Exception as e:
            logging.warning(f"Processing failed: {e}")
            self.download_status = "failed_processing"
            self.error_message = str(e)
        else:
            return result

    @abstractmethod
    def _handle(
        self,
        file_path: str,
        dest_file_path: str,
    ) -> ProcessingResult: ...

    @classmethod
    def _mkdir(cls, folder: str) -> Path:
        folder: Path = Path(folder)
        folder.mkdir(parents=True, exist_ok=True)
        return folder


class PDFHandler(DocumentHandler):
    def _handle(
        self,
        file_path: str,
        dest_file_path: str,
    ) -> ProcessingResult:
        pdf_reader = PdfReader(file_path)
        metadata: DocumentInformation = pdf_reader.metadata
        text: str = "\n".join(page.extract_text() for page in pdf_reader.pages)

        return ProcessingResult(
            detected_language=langdetect.detect(text),
            document_page_count=len(pdf_reader.pages),
            author=metadata.author,
            creation_date=utils.parse_date(metadata.creation_date_raw),
        )


class DocXHandler(DocumentHandler):
    def _handle(
        self,
        file_path: str,
        dest_file_path: str,
    ) -> ProcessingResult:
        document = docx.Document(file_path)
        metadata = document.core_properties
        text: str = "\n".join(paragraph.text for paragraph in document.paragraphs)

        creation_date = metadata.created
        self.metadata = {
            "document_page_count": len(document.paragraphs),
            "author": metadata.author,
            "creation_date": creation_date.strftime("%Y-%m-%d %H:%M:%S") if creation_date else None,
            "language": langdetect.detect(text),
        }


class XLSXHandler(DocumentHandler):
    def _handle(
        self,
        file_path: str,
        dest_file_path: str,
    ) -> ProcessingResult:
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        metadata = workbook.properties
        text: str = "\n".join(
            " ".join(str(cell) for cell in row if cell is not None)
            for sheet in workbook.worksheets
            for row in sheet.iter_rows(values_only=True)
        )

        creation_date = metadata.created
        self.metadata = {
            "author": metadata.creator,
            "creation_date": creation_date.strftime("%Y-%m-%d %H:%M:%S") if creation_date else None,
            "language": langdetect.detect(text),
        }


class HTMLHandler(DocumentHandler):
    def _handle(
        self,
        file_path: str,
        dest_file_path: str,
    ) -> ProcessingResult:
        with open(file_path, "r") as input_file:
            soup = BeautifulSoup(input_file, "html.parser")
            text = soup.text

            self.metadata = {"language": langdetect.detect(soup.text)}


class HandlerFactory:
    _map = {
        ".pdf": PDFHandler,
        ".docx": DocXHandler,
        ".xlsx": XLSXHandler,
    }

    @classmethod
    def get_handler(cls, file_path: Path, output_dir: Path) -> DocumentHandler:
        ext = file_path.suffix.lower()
        handler_cls = cls._map.get(ext)
        if not handler_cls:
            raise ValueError(f"No handler for extension '{ext}'")
        return handler_cls(output_dir=output_dir)
