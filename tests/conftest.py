from unittest.mock import MagicMock
from typing import Any
from datetime import (
    datetime,
    timedelta,
)
import random
import uuid
import string
import io

import pytest
import docx
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.pagesizes import A4

from domain.document.service import DocumentService
from domain.document.schemas import (
    File,
    Document,
    DocumentStatus,
)
from domain.chat.service import (
    RAGService,
)
from domain.chat.repositories import (
    ChatSessionRepository,
    ChatMessageRepository,
)
from domain.workspace.repositories import WorkspaceRepository
from services import (
    RawStorage,
    VectorStore,
)
from domain.embedding import (
    VectorMetadata,
    Vector,
)
from domain.database.uow import UnitOfWork
from domain.embedding import EmbeddingModel
from domain.text_splitter import TextSplitter


class ValueGenerator:
    @classmethod
    def float_list(
        cls,
        min_f: int = 0,
        max_f: int = 1,
        n_values: int | None = None,
        exclude: set[float] | None = None,
    ) -> list[float]:
        if exclude is None:
            exclude = {}
        n_values = n_values or 10
        floats: list[float] = []

        def random_float():
            return random.randrange(min_f, max_f + 1) * random.random()

        for _ in range(n_values):
            while (value := random_float()) in exclude:
                pass
            floats.append(value)

        return floats

    @classmethod
    def word(cls, n: int = 10) -> str:
        return "".join(random.choices(string.ascii_letters, k=n))

    @classmethod
    def text(cls, n: int = 10) -> str:
        return "".join(
            random.choices(
                string.ascii_letters + string.digits + string.punctuation + " ", k=n
            )
        )

    @classmethod
    def integer(cls, n: int = 10) -> int:
        return int("".join(random.choices(string.digits, k=n)))

    @classmethod
    def uuid(cls, length: int | None = None) -> str:
        if length:
            return str(uuid.uuid4())[:length]
        return str(uuid.uuid4())

    @classmethod
    def datetime(
        cls,
        start: datetime = datetime(2020, 1, 1, 0, 0, 0),
        end: datetime = datetime(2025, 1, 1, 0, 0, 0),
    ):
        delta = end - start
        random_seconds = random.randint(0, int(delta.total_seconds()))
        return start + timedelta(seconds=random_seconds)

    @classmethod
    def path(cls, sub_directories: int = 0) -> str:
        return f"{'/'.join([cls.word() for _ in range(sub_directories + 1)])}/"

    @classmethod
    def vector(cls, n_values: int = 384) -> Vector:
        return Vector(
            values=cls.float_list(-1, 1, n_values, {0}),
            metadata=VectorMetadata(
                document_id=cls.uuid(),
                workspace_id=cls.uuid(),
                document_name=cls.text(),
                page_start=cls.integer(),
                page_end=cls.integer(),
                text=cls.text(),
            ),
        )

    @classmethod
    def vectors(cls, n_vectors: int = 10, n_values: int = 384) -> list[Vector]:
        return [cls.vector(n_values) for _ in range(n_vectors)]

    @classmethod
    def float_vector(cls, n_values: int = 384) -> list[float]:
        return cls.float_list(-1, 1, n_values, {0})

    @classmethod
    def chunks(cls, n_values: int) -> list[str]:
        return [cls.word() for _ in range(n_values)]

    @classmethod
    def document_metadata(cls) -> Document:
        return Document(
            document_id=cls.uuid(),
            workspace_id=cls.uuid(),
            document_name=cls.text(),
            media_type="application/pdf",
            detected_language="en",
            document_page_count=cls.integer(),
            author=cls.text(),
            creation_date=cls.datetime(),
            raw_storage_path=f"{cls.path()}{cls.word()}.pdf",
            file_size_bytes=cls.integer(),
            ingested_at=cls.datetime(),
            status=DocumentStatus.success,
        )

    @classmethod
    def pdf(
        cls,
        path: str,
        target_bytes: int = 1_000_000,
        min_pages: int = 1,
        max_pages: int = 1000,
        min_page_text_num: int = 500,
        max_page_text_num: int = 5000,
        min_line_len: int = 20,
        max_line_len: int = 100,
    ) -> None:
        buffer = io.BytesIO()
        document = Canvas(buffer, pagesize=A4)
        num_pages = random.randint(min_pages, max_pages)

        for _ in range(num_pages):
            text: str = cls.text(random.randint(min_page_text_num, max_page_text_num))
            text_obj = document.beginText(40, 800)

            start: int = 0
            end: int = len(text)
            while start < end:
                line_len = random.randint(min_line_len, max_line_len)
                text_obj.textLine(text[start : start + line_len])
                start += line_len

            document.drawText(text_obj)
            document.showPage()

        document.save()
        data = buffer.getvalue()

        if len(data) < target_bytes:
            data += b"\n" * (target_bytes - len(data))

        with open(path, "wb") as f:
            f.write(data)

    @classmethod
    def docx(
        cls,
        path: str,
        target_bytes: int,
    ) -> None:
        buffer = io.BytesIO()
        document = docx.Document()
        num_pages: int = int(target_bytes / 1_000_000)
        page_text_num: int = int(target_bytes / num_pages)

        for page_num in range(num_pages):
            text: str = cls.text(page_text_num)
            document.add_paragraph(text)
            if page_num < num_pages - 1:
                document.add_page_break()

        document.save(buffer)
        data = buffer.getvalue()

        for _ in range(int((target_bytes - len(data)) / page_text_num)):
            text: str = cls.text(page_text_num)
            document.add_paragraph(text)

        document.save(path)


@pytest.fixture
def tmp_document(tmp_path) -> Any:
    _map: dict[int, tuple] = {
        0: (ValueGenerator.pdf, ".pdf"),
        1: (ValueGenerator.docx, ".docx"),
    }

    def _make(
        target_bytes: int = 1_000_000, doc_type: str | None = None
    ) -> tuple[bytes, str, str]:
        if doc_type:
            if doc_type == ".pdf":
                func, file_extension = _map[0]
            elif doc_type == ".docx":
                func, file_extension = _map[1]
            else:
                raise NotImplementedError()
        else:
            func, file_extension = _map[random.randint(0, len(_map) - 1)]
        path = tmp_path / f"document_{target_bytes}{file_extension}"
        func(path, target_bytes)
        with open(path, "rb") as file:
            file_bytes: bytes = file.read()
        return file_bytes, str(path), file_extension

    return _make


def assert_any_exception(exctype, excinfo) -> None:
    assert excinfo.type is exctype
    assert excinfo.value.message == exctype.message
    assert excinfo.value.error_code == exctype.error_code
    assert excinfo.value.status_code == exctype.status_code


@pytest.fixture
def mock_raw_storage(mocker) -> MagicMock:
    return mocker.create_autospec(RawStorage, instance=True)


@pytest.fixture
def mock_vector_store(mocker) -> MagicMock:
    return mocker.create_autospec(VectorStore, instance=True)


@pytest.fixture
def mock_embedding_model(mocker) -> MagicMock:
    return mocker.create_autospec(EmbeddingModel, instance=True)


@pytest.fixture
def mock_text_splitter(mocker) -> MagicMock:
    return mocker.create_autospec(TextSplitter, instance=True)


@pytest.fixture
def mock_document_service(mocker) -> MagicMock:
    return mocker.create_autospec(DocumentService, instance=True)


@pytest.fixture
def mock_file_scheme(mocker) -> MagicMock:
    return mocker.create_autospec(File, instance=True)


@pytest.fixture
def mock_chat_session_repository(mocker) -> MagicMock:
    return mocker.create_autospec(ChatSessionRepository, instance=True)


@pytest.fixture
def mock_chat_message_repository(mocker) -> MagicMock:
    return mocker.create_autospec(ChatMessageRepository, instance=True)


@pytest.fixture
def mock_rag_service(mocker) -> MagicMock:
    return mocker.create_autospec(RAGService, instance=True)


@pytest.fixture
def mock_workspace_repository(mocker) -> MagicMock:
    return mocker.create_autospec(WorkspaceRepository, instance=True)


@pytest.fixture
def mock_uow(mocker) -> MagicMock:
    return mocker.create_autospec(UnitOfWork, instance=True)
