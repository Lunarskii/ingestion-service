from datetime import (
    datetime,
    timedelta,
)
import random
import io
import uuid
import string

from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.pagesizes import A4
import docx

from app.domain.document.schemas import (
    File,
    Document,
    DocumentDTO,
    DocumentStatus,
)
from app.types import VectorPayload, Vector
from app.domain.extraction import Page
from app.domain.text_splitter import (
    Chunk,
    PageSpan,
)


class DocumentGenerator:
    @classmethod
    def document_dto(cls, n: int = 1) -> DocumentDTO | list[DocumentDTO]:
        assert n > 0

        documents: list[DocumentDTO] = [
            DocumentDTO(
                id=ValueGenerator.uuid(),
                workspace_id=ValueGenerator.uuid(),
                source_id=ValueGenerator.text(),
                run_id=random.choice((None, ValueGenerator.uuid())),
                trace_id=ValueGenerator.uuid(),
                sha256=ValueGenerator.word(),
                raw_url=random.choice((None, ValueGenerator.text())),
                title=ValueGenerator.text(),
                media_type=ValueGenerator.text(),
                detected_language=random.choice((None, ValueGenerator.word())),
                page_count=random.choice((None, ValueGenerator.integer())),
                author=random.choice((None, ValueGenerator.text())),
                creation_date=random.choice((None, ValueGenerator.datetime())),
                raw_storage_path=f"{ValueGenerator.path()}{ValueGenerator.word()}.pdf",
                silver_storage_path=f"{ValueGenerator.path()}{ValueGenerator.word()}.json",
                size_bytes=ValueGenerator.integer(),
                fetched_at=ValueGenerator.datetime(),
                stored_at=ValueGenerator.datetime(),
                ingested_at=ValueGenerator.datetime(),
                status=DocumentStatus.success,
                error_message=random.choice((None, ValueGenerator.text())),
            )
            for _ in range(n)
        ]
        if n == 1:
            return documents[0]
        return documents

    @classmethod
    def document(cls, n: int = 1) -> Document | list[Document]:
        assert n > 0

        documents: list[Document] = [
            Document(
                id=ValueGenerator.uuid(),
                workspace_id=ValueGenerator.uuid(),
                source_id=ValueGenerator.text(),
                run_id=random.choice((None, ValueGenerator.uuid())),
                trace_id=ValueGenerator.uuid(),
                sha256=ValueGenerator.word(),
                raw_url=random.choice((None, ValueGenerator.text())),
                title=ValueGenerator.text(),
                media_type=ValueGenerator.text(),
                detected_language=random.choice((None, ValueGenerator.word())),
                page_count=random.choice((None, ValueGenerator.integer())),
                author=random.choice((None, ValueGenerator.text())),
                creation_date=random.choice((None, ValueGenerator.datetime())),
                raw_storage_path=f"{ValueGenerator.path()}{ValueGenerator.word()}.pdf",
                silver_storage_path=f"{ValueGenerator.path()}{ValueGenerator.word()}.json",
                size_bytes=ValueGenerator.integer(),
                fetched_at=ValueGenerator.datetime(),
                stored_at=ValueGenerator.datetime(),
                ingested_at=ValueGenerator.datetime(),
                status=DocumentStatus.success,
                error_message=random.choice((None, ValueGenerator.text())),
            )
            for _ in range(n)
        ]
        if n == 1:
            return documents[0]
        return documents

    @classmethod
    def pdf(
        cls,
        path: str,
        target_bytes: int = 1_000_000,
        min_pages: int = 1,
        max_pages: int = 1000,
        min_page_text_num: int = 500,
        max_page_text_num: int = 5000,
        min_line_len: int = 20,
        max_line_len: int = 100,
    ) -> None:
        buffer = io.BytesIO()
        document = Canvas(buffer, pagesize=A4)
        num_pages = random.randint(min_pages, max_pages)

        for _ in range(num_pages):
            text: str = ValueGenerator.text(random.randint(min_page_text_num, max_page_text_num))
            text_obj = document.beginText(40, 800)

            start: int = 0
            end: int = len(text)
            while start < end:
                line_len = random.randint(min_line_len, max_line_len)
                text_obj.textLine(text[start : start + line_len])
                start += line_len

            document.drawText(text_obj)
            document.showPage()

        document.save()
        data = buffer.getvalue()

        if len(data) < target_bytes:
            data += b"\n" * (target_bytes - len(data))

        with open(path, "wb") as f:
            f.write(data)

    @classmethod
    def docx(
        cls,
        path: str,
        target_bytes: int,
    ) -> None:
        buffer = io.BytesIO()
        document = docx.Document()
        num_pages: int = int(target_bytes / 1_000_000)
        page_text_num: int = int(target_bytes / num_pages)

        for page_num in range(num_pages):
            text: str = ValueGenerator.text(page_text_num)
            document.add_paragraph(text)
            if page_num < num_pages - 1:
                document.add_page_break()

        document.save(buffer)
        data = buffer.getvalue()

        for _ in range(int((target_bytes - len(data)) / page_text_num)):
            text: str = ValueGenerator.text(page_text_num)
            document.add_paragraph(text)

        document.save(path)


class ValueGenerator:
    @classmethod
    def float_list(
        cls,
        min_f: int = 0,
        max_f: int = 1,
        n_values: int | None = None,
        exclude: set[float] | None = None,
    ) -> list[float]:
        if exclude is None:
            exclude = {}
        n_values = n_values or 10
        floats: list[float] = []

        def random_float():
            return random.randrange(min_f, max_f + 1) * random.random()

        for _ in range(n_values):
            while (value := random_float()) in exclude:
                pass
            floats.append(value)

        return floats

    @classmethod
    def word(cls, n: int = 10) -> str:
        return "".join(random.choices(string.ascii_letters, k=n))

    @classmethod
    def text(cls, n: int = 10) -> str:
        return "".join(
            random.choices(
                string.ascii_letters + string.digits + string.punctuation + " ", k=n
            )
        )

    @classmethod
    def integer(cls, n: int = 10) -> int:
        return int("".join(random.choices(string.digits, k=n)))

    @classmethod
    def uuid(cls, length: int | None = None) -> str:
        if length:
            return str(uuid.uuid4())[:length]
        return str(uuid.uuid4())

    @classmethod
    def datetime(
        cls,
        start: datetime = datetime(2020, 1, 1, 0, 0, 0),
        end: datetime = datetime(2025, 1, 1, 0, 0, 0),
    ):
        delta = end - start
        random_seconds = random.randint(0, int(delta.total_seconds()))
        return start + timedelta(seconds=random_seconds)

    @classmethod
    def bytes(cls, length: int = 10) -> bytes:
        return random.randbytes(length)

    @classmethod
    def path(cls, sub_directories: int = 0) -> str:
        return f"{'/'.join([cls.word() for _ in range(sub_directories + 1)])}/"

    # @classmethod
    # def vector(cls, n_values: int = 384) -> Vector:
    #     return Vector(
    #         values=cls.float_list(-1, 1, n_values, {0}),
    #         metadata=VectorMetadata(
    #             document_id=cls.uuid(),
    #             workspace_id=cls.uuid(),
    #             document_name=cls.text(),
    #             page_start=cls.integer(),
    #             page_end=cls.integer(),
    #             text=cls.text(),
    #         ),
    #     )
    #
    # @classmethod
    # def vectors(cls, n_vectors: int = 10, n_values: int = 384) -> list[Vector]:
    #     return [cls.vector(n_values) for _ in range(n_vectors)]

    @classmethod
    def float_vector(cls, n_values: int = 384) -> list[float]:
        return cls.float_list(-1, 1, n_values, {0})

    @classmethod
    def page(cls, n: int = 1) -> Page | list[Page]:
        assert n > 0

        pages: list[Page] = [
            Page(
                num=cls.integer(),
                text=cls.text(),
            )
            for _ in range(n)
        ]
        if n == 1:
            return pages[0]
        return pages

    @classmethod
    def chunk(cls, n: int = 1) -> Chunk | list[Chunk]:
        assert n > 0

        chunks: list[Chunk] = [
            Chunk(
                text=cls.text(),
                page_spans=[
                    PageSpan(
                        text=cls.text(),
                        page_num=page_num,
                        chunk_start_on_page=page_num,
                        chunk_end_on_page=page_num,
                    )
                    for page_num in range(random.randint(1, 2))
                ],
            )
            for _ in range(n)
        ]
        if n == 1:
            return chunks[0]
        return chunks

    @classmethod
    def vector(
        cls,
        n: int = 1,
        n_values: int = 384,
        chunks: list[Chunk] | None = None,
        document: DocumentDTO | None = None,
    ) -> Vector | list[Vector]:
        assert n > 0 or chunks is not None

        if chunks:
            n = len(chunks)
            vectors: list[Vector] = [
                Vector(
                    values=cls.float_list(-1, 1, n_values, {0}),
                    payload=VectorPayload(
                        document_id=document.id if document else cls.uuid(),
                        workspace_id=document.workspace_id if document else cls.uuid(),
                        document_name=document.title if document else cls.word(),
                        page_start=chunk.page_spans[0].page_num,
                        page_end=chunk.page_spans[-1].page_num,
                        text=chunk.text,
                    ),
                )
                for chunk in chunks
            ]
        else:
            vectors: list[Vector] = [
                Vector(
                    values=cls.float_list(-1, 1, n_values, {0}),
                    payload=VectorPayload(
                        document_id=cls.uuid(),
                        workspace_id=cls.uuid(),
                        document_name=cls.word(),
                        page_start=cls.integer(),
                        page_end=cls.integer(),
                        text=cls.text(),
                    ),
                )
                for _ in range(n)
            ]
        if n == 1:
            return vectors[0]
        return vectors
